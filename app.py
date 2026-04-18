import os
from datetime import datetime
from flask import Flask, request, jsonify, send_from_directory, redirect
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
import requests

app = Flask(__name__, static_folder='static')
CORS(app, resources={r"/exam/*": {"origins": "*"}})

basedir = os.path.abspath(os.path.dirname(__file__))
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(basedir, 'exam.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = os.path.join(basedir, 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)

WECHAT_WORK_API_URL = "https://qyapi.weixin.qq.com/cgi-bin/message/send"
WECHAT_WORK_GETTOKEN_URL = "https://qyapi.weixin.qq.com/cgi-bin/gettoken"

class Exam(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.now)
    questions = db.relationship('Question', backref='exam', lazy=True, cascade='all, delete-orphan')

    def to_dict(self):
        return {
            'id': self.id,
            'title': self.title,
            'description': self.description,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'questions': [q.to_dict() for q in self.questions]
        }

class Question(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    exam_id = db.Column(db.Integer, db.ForeignKey('exam.id'), nullable=False)
    question_text = db.Column(db.Text, nullable=False)
    question_type = db.Column(db.String(50), nullable=False)
    options = db.Column(db.Text)
    correct_answer = db.Column(db.String(500), nullable=False)
    score = db.Column(db.Integer, default=10)

    def to_dict(self):
        return {
            'id': self.id,
            'question_text': self.question_text,
            'question_type': self.question_type,
            'options': self.options.split('|') if self.options else [],
            'correct_answer': self.correct_answer,
            'score': self.score
        }

class Submission(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    exam_id = db.Column(db.Integer, db.ForeignKey('exam.id'), nullable=False)
    student_name = db.Column(db.String(100), nullable=False)
    employee_id = db.Column(db.String(100), nullable=True)
    answers = db.Column(db.Text)
    score = db.Column(db.Float, default=0)
    submitted_at = db.Column(db.DateTime, default=datetime.now)
    graded = db.Column(db.Boolean, default=False)

    def to_dict(self, include_questions=False):
        result = {
            'id': self.id,
            'exam_id': self.exam_id,
            'student_name': self.student_name,
            'employee_id': self.employee_id,
            'answers': self.answers.split('|') if self.answers else [],
            'score': self.score,
            'submitted_at': self.submitted_at.isoformat() if self.submitted_at else None,
            'graded': self.graded
        }
        if include_questions:
            exam = Exam.query.get(self.exam_id)
            if exam:
                result['exam_title'] = exam.title
                result['questions'] = [q.to_dict() for q in exam.questions]
            else:
                result['exam_title'] = '试卷已删除'
                result['questions'] = []
        return result

class WebhookConfig(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    webhook_url = db.Column(db.String(500), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.now)

    def to_dict(self):
        return {
            'id': self.id,
            'name': self.name,
            'webhook_url': self.webhook_url,
            'created_at': self.created_at.isoformat() if self.created_at else None
        }

with app.app_context():
    db.create_all()

@app.route('/exam/')
def index():
    return send_from_directory('static', 'index.html')

@app.route('/exam/exam/<int:exam_id>')
def exam_page(exam_id):
    return send_from_directory('static', 'exam.html')

@app.route('/exam/api/exams', methods=['GET'])
def get_exams():
    exams = Exam.query.order_by(Exam.created_at.desc()).all()
    return jsonify([e.to_dict() for e in exams])

@app.route('/exam/api/exams/<int:exam_id>', methods=['GET'])
def get_exam(exam_id):
    exam = Exam.query.get_or_404(exam_id)
    return jsonify(exam.to_dict())

@app.route('/exam/api/exams', methods=['POST'])
def create_exam():
    data = request.json

    exam = Exam(
        title=data.get('title'),
        description=data.get('description', '')
    )
    db.session.add(exam)
    db.session.flush()

    for q in data.get('questions', []):
        question = Question(
            exam_id=exam.id,
            question_text=q.get('question_text'),
            question_type=q.get('question_type'),
            options='|'.join(q.get('options', [])) if q.get('options') else '',
            correct_answer=q.get('correct_answer'),
            score=q.get('score', 10)
        )
        db.session.add(question)

    db.session.commit()
    return jsonify(exam.to_dict()), 201

@app.route('/exam/api/exams/<int:exam_id>', methods=['PUT'])
def update_exam(exam_id):
    exam = Exam.query.get_or_404(exam_id)
    data = request.json

    exam.title = data.get('title', exam.title)
    exam.description = data.get('description', exam.description)

    if 'questions' in data:
        Question.query.filter_by(exam_id=exam_id).delete()
        for q in data.get('questions', []):
            question = Question(
                exam_id=exam.id,
                question_text=q.get('question_text'),
                question_type=q.get('question_type'),
                options='|'.join(q.get('options', [])) if q.get('options') else '',
                correct_answer=q.get('correct_answer'),
                score=q.get('score', 10)
            )
            db.session.add(question)

    db.session.commit()
    return jsonify(exam.to_dict())

@app.route('/exam/api/exams/<int:exam_id>', methods=['DELETE'])
def delete_exam(exam_id):
    exam = Exam.query.get_or_404(exam_id)
    db.session.delete(exam)
    db.session.commit()
    return jsonify({'message': '试卷已删除'})

@app.route('/exam/api/submit', methods=['POST'])
def submit_exam():
    data = request.json
    exam_id = data.get('exam_id')
    student_name = data.get('student_name')
    employee_id = data.get('employee_id', '')
    answers = data.get('answers', [])

    exam = Exam.query.get_or_404(exam_id)
    submission = Submission(
        exam_id=exam_id,
        student_name=student_name,
        employee_id=employee_id,
        answers='|'.join(str(a) for a in answers)
    )
    db.session.add(submission)
    db.session.flush()

    score = grade_exam(exam, answers)
    submission.score = score
    submission.graded = True

    db.session.commit()
    return jsonify(submission.to_dict())

def grade_exam(exam, answers):
    total_score = 0
    questions = exam.questions

    for i, question in enumerate(questions):
        if i < len(answers):
            user_answer = str(answers[i]).strip()
            correct_answer = str(question.correct_answer).strip()

            if question.question_type == 'single':
                if user_answer.upper() == correct_answer.upper():
                    total_score += question.score
            elif question.question_type == 'multiple':
                user_set = set(u.strip().upper() for u in user_answer.split(','))
                correct_set = set(c.strip().upper() for c in correct_answer.split(','))
                if user_set == correct_set:
                    total_score += question.score
            elif question.question_type == 'truefalse':
                if user_answer.upper() == correct_answer.upper():
                    total_score += question.score
            elif question.question_type == 'fillblank':
                if user_answer.upper() == correct_answer.upper():
                    total_score += question.score
            elif question.question_type == 'essay':
                score = 0
                if user_answer and correct_answer:
                    keywords = [k.strip() for k in correct_answer.split('|') if k.strip()]
                    if not keywords:
                        keywords = [correct_answer]
                    
                    matched_keywords = 0
                    for keyword in keywords:
                        if keyword.lower() in user_answer.lower():
                            matched_keywords += 1
                    
                    if matched_keywords > 0:
                        score = int(question.score * (matched_keywords / len(keywords)))
                total_score += score

    return total_score

@app.route('/exam/api/submissions/<int:exam_id>', methods=['GET'])
def get_submissions(exam_id):
    submissions = Submission.query.filter_by(exam_id=exam_id).order_by(Submission.submitted_at.desc()).all()
    return jsonify([s.to_dict() for s in submissions])

@app.route('/exam/api/send-to-wechat', methods=['POST'])
def send_to_wechat():
    data = request.json
    exam_id = data.get('exam_id')
    webhook_url = data.get('webhook_url')

    exam = Exam.query.get_or_404(exam_id)

    if not webhook_url:
        return jsonify({'error': '请提供企业微信机器人Webhook地址'}), 400

    exam_url = f"{request.host_url}exam/exam/{exam_id}"

    message = {
        "msgtype": "news",
        "news": {
            "articles": [
                {
                    "title": f"📝 {exam.title}",
                    "description": f"{exam.description or '请点击下方链接查看试卷并完成答题'}",
                    "url": exam_url,
                    "picurl": ""
                }
            ]
        }
    }

    try:
        response = requests.post(webhook_url, json=message, timeout=10)
        result = response.json()

        if result.get('errcode') == 0:
            return jsonify({
                'message': '试卷已成功发送到企业微信',
                'exam_url': exam_url
            })
        else:
            return jsonify({'error': result.get('errmsg', '发送失败')}), 400
    except requests.exceptions.RequestException as e:
        return jsonify({'error': f'发送失败: {str(e)}'}), 400

@app.route('/exam/api/webhooks', methods=['GET'])
def get_webhooks():
    webhooks = WebhookConfig.query.order_by(WebhookConfig.created_at.desc()).all()
    return jsonify([w.to_dict() for w in webhooks])

@app.route('/exam/api/webhooks', methods=['POST'])
def create_webhook():
    data = request.json
    webhook = WebhookConfig(
        name=data.get('name', '未命名'),
        webhook_url=data.get('webhook_url')
    )
    db.session.add(webhook)
    db.session.commit()
    return jsonify(webhook.to_dict()), 201

@app.route('/exam/api/webhooks/<int:webhook_id>', methods=['DELETE'])
def delete_webhook(webhook_id):
    webhook = WebhookConfig.query.get_or_404(webhook_id)
    db.session.delete(webhook)
    db.session.commit()
    return jsonify({'message': 'Webhook已删除'})

@app.route('/exam/api/submission/<int:submission_id>', methods=['GET'])
def get_submission(submission_id):
    submission = Submission.query.get_or_404(submission_id)
    return jsonify(submission.to_dict(include_questions=True))

@app.route('/exam/api/submission/<int:submission_id>', methods=['DELETE'])
def delete_submission(submission_id):
    submission = Submission.query.get_or_404(submission_id)
    db.session.delete(submission)
    db.session.commit()
    return jsonify({'message': '提交记录已删除'})

@app.route('/exam/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': '没有文件上传'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '没有选择文件'}), 400

    filename = file.filename
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    try:
        questions = []
        if filename.endswith('.xlsx') or filename.endswith('.xls'):
            questions = parse_excel(filepath)
        elif filename.endswith('.docx'):
            questions = parse_word(filepath)
        else:
            return jsonify({'error': '不支持的文件格式，请上传Excel或Word文件'}), 400

        return jsonify({
            'message': '文件解析成功',
            'questions': questions,
            'filename': filename
        })
    except Exception as e:
        return jsonify({'error': f'解析文件失败: {str(e)}'}), 400
    finally:
        if os.path.exists(filepath):
            os.remove(filepath)

def parse_excel(filepath):
    import openpyxl

    questions = []
    wb = openpyxl.load_workbook(filepath)

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]

        for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            if not row or not any(cell for cell in row if cell):
                continue

            cells = [str(cell) if cell else '' for cell in row]

            if not cells[0] or not cells[0].strip():
                continue

            q_type = 'single'
            q_text = ''
            q_options = []
            q_answer = ''
            q_score = 10

            if len(cells) >= 1 and cells[0]:
                q_text = cells[0].strip()

            if len(cells) >= 2 and cells[1]:
                type_text = cells[1].strip().lower()
                if '单' in type_text:
                    q_type = 'single'
                elif '多' in type_text:
                    q_type = 'multiple'
                elif '判断' in type_text:
                    q_type = 'truefalse'
                elif '填空' in type_text:
                    q_type = 'fillblank'

            if q_type == 'truefalse':
                q_options = ['正确', '错误']
                if len(cells) >= 3 and cells[2]:
                    answer_text = cells[2].strip().upper()
                    if '错' in answer_text or 'B' in answer_text or 'FALSE' in answer_text:
                        q_answer = 'FALSE'
                    else:
                        q_answer = 'TRUE'
            elif q_type == 'fillblank':
                q_options = []
                if len(cells) >= 3 and cells[2]:
                    q_answer = cells[2].strip()
            else:
                option_idx = 2
                while option_idx < len(cells) - 2 and len(q_options) < 4:
                    if cells[option_idx] and cells[option_idx].strip():
                        q_options.append(cells[option_idx].strip())
                    option_idx += 1

                if len(cells) >= 3:
                    last_option_idx = 2 + len(q_options)
                    if last_option_idx < len(cells) and cells[last_option_idx]:
                        q_answer = cells[last_option_idx].strip()
                        if q_type == 'multiple' and ',' not in q_answer:
                            q_answer = q_answer.upper().replace(' ', ',')

                if len(cells) >= 3:
                    last_idx = len(cells) - 1
                    if last_idx > 2 + len(q_options):
                        try:
                            q_score = int(cells[last_idx])
                        except:
                            pass

            if q_text and q_answer:
                questions.append({
                    'question_text': q_text,
                    'question_type': q_type,
                    'options': q_options,
                    'correct_answer': q_answer,
                    'score': q_score
                })

    return questions

def parse_word(filepath):
    from docx import Document

    questions = []
    doc = Document(filepath)

    current_question = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if text[0].isdigit() and '.' in text[:5]:
            if current_question and current_question['question_text']:
                questions.append(current_question)

            current_question = {
                'question_text': '',
                'question_type': 'single',
                'options': [],
                'correct_answer': '',
                'score': 10
            }

            parts = text.split('.', 1)
            if len(parts) > 1:
                current_question['question_text'] = parts[1].strip()
            else:
                current_question['question_text'] = text

            text_lower = text.lower()
            if '单选' in text or '单选题' in text:
                current_question['question_type'] = 'single'
            elif '多选' in text or '多选题' in text:
                current_question['question_type'] = 'multiple'
            elif '判断' in text or '判断题' in text:
                current_question['question_type'] = 'truefalse'
                current_question['options'] = ['正确', '错误']
            elif '填空' in text or '填空题' in text:
                current_question['question_type'] = 'fillblank'

    if current_question and current_question['question_text']:
        questions.append(current_question)

    for para in doc.paragraphs:
        text = para.text.strip()
        if len(text) == 1 and text.upper() in ['A', 'B', 'C', 'D']:
            if questions and questions[-1]['question_type'] == 'single':
                for i, opt in enumerate(['A', 'B', 'C', 'D']):
                    if text.upper() == opt:
                        next_paras = []
                        for j in range(doc.paragraphs.index(para) + 1, min(doc.paragraphs.index(para) + 5, len(doc.paragraphs))):
                            next_text = doc.paragraphs[j].text.strip()
                            if next_text and not next_text[0].isdigit():
                                next_paras.append(next_text)
                            else:
                                break

                        if next_paras:
                            full_option = text + '. ' + ' '.join(next_paras)
                            if len(questions[-1]['options']) <= i:
                                questions[-1]['options'].append(full_option)
                            break

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if cells and cells[0] and cells[0][0].isdigit():
                q_text = ''
                q_type = 'single'
                q_options = []
                q_answer = ''
                q_score = 10

                parts = cells[0].split('.', 1)
                if len(parts) > 1:
                    q_text = parts[1].strip()
                else:
                    q_text = cells[0]

                if len(cells) >= 2:
                    type_text = cells[1].lower()
                    if '单' in type_text:
                        q_type = 'single'
                    elif '多' in type_text:
                        q_type = 'multiple'
                    elif '判断' in type_text:
                        q_type = 'truefalse'
                    elif '填空' in type_text:
                        q_type = 'fillblank'

                if q_type == 'truefalse':
                    q_options = ['正确', '错误']
                    if len(cells) >= 3:
                        answer_text = cells[2].strip().upper()
                        q_answer = 'FALSE' if ('错' in answer_text or 'B' in answer_text) else 'TRUE'
                elif q_type == 'fillblank':
                    q_options = []
                    if len(cells) >= 3:
                        q_answer = cells[2].strip()
                else:
                    if len(cells) >= 3 and cells[2]:
                        q_options.append(cells[2])
                    if len(cells) >= 4 and cells[3]:
                        q_options.append(cells[3])
                    if len(cells) >= 5 and cells[4]:
                        q_options.append(cells[4])
                    if len(cells) >= 6 and cells[5]:
                        q_options.append(cells[5])

                    if len(cells) >= 7 and cells[6]:
                        q_answer = cells[6].strip()

                if len(cells) >= 8 and cells[7]:
                    try:
                        q_score = int(cells[7])
                    except:
                        q_score = 10

                if q_text:
                    questions.append({
                        'question_text': q_text,
                        'question_type': q_type,
                        'options': q_options,
                        'correct_answer': q_answer,
                        'score': q_score
                    })

    return questions

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8000))
    app.run(host='0.0.0.0', port=port)